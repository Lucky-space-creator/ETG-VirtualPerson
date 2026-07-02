from docx import Document
import io


class DocxParser:
    """DOCX 文档解析器，使用 python-docx"""

    def parse(self, content: bytes, filename: str) -> dict:
        try:
            stream = io.BytesIO(content)
            doc = Document(stream)

            text_parts = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    text_parts.append(text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)
        except Exception as e:
            raise RuntimeError(f"DOCX解析失败: {str(e)}")

        return {
            "text": full_text,
            "page_count": 1,
            "char_count": len(full_text),
        }
